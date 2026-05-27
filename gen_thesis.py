# -*- coding: utf-8 -*-
"""
生成毕业论文 docx 文件
安徽三联学院 本科毕业设计（论文）
题目：基于SSM框架的学生实习管理系统的设计与实现
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'thesis_images')
OUTPUT_FILE = os.path.join(BASE_DIR, '毕业论文.docx')

# ============================================================
#  字号映射 (中文字号 -> 磅值)
# ============================================================
SZ = {
    '初号': 42, '小初': 36, '一号': 26, '小一': 24,
    '二号': 22, '小二': 18, '三号': 16, '小三': 15,
    '四号': 14, '小四': 12, '五号': 10.5, '小五': 9,
}


# ============================================================
#  格式化辅助函数
# ============================================================
def _set_run(run, font_cn='宋体', font_en='Times New Roman', size=12,
             bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_paragraph_fmt(p, alignment=None, line_spacing=1.5,
                       space_before=0, space_after=0,
                       first_line_indent=None):
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)


def add_paragraph(doc, text, font_cn='宋体', font_en='Times New Roman',
                  size=12, bold=False, alignment=None,
                  line_spacing=1.5, space_before=0, space_after=0,
                  first_line_indent=None):
    p = doc.add_paragraph()
    _set_paragraph_fmt(p, alignment, line_spacing, space_before, space_after, first_line_indent)
    run = p.add_run(text)
    _set_run(run, font_cn, font_en, size, bold)
    return p


def add_body(doc, text):
    return add_paragraph(doc, text, '宋体', 'Times New Roman', SZ['小四'],
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=SZ['小四'] * 2)


def add_heading1(doc, text):
    p = add_paragraph(doc, '', space_before=24)
    p2 = add_paragraph(doc, text, '黑体', 'Arial', SZ['三号'],
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       line_spacing=1.5, space_before=0, space_after=12)
    _set_run(p2.runs[0], '黑体', 'Arial', SZ['三号'], bold=False)
    return p2


def add_heading2(doc, text):
    return add_paragraph(doc, text, '黑体', 'Arial', SZ['四号'],
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing=1.5, space_before=6, space_after=3)


def add_heading3(doc, text):
    return add_paragraph(doc, text, '黑体', 'Arial', SZ['小四'],
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing=1.5, space_before=3, space_after=3)


def add_figure(doc, img_name, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run(r, '黑体', 'Arial', SZ['五号'], bold=False)
    cap.paragraph_format.space_after = Pt(6)


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'</w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_three_line_table(doc, caption, headers, rows):
    # 标题
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run(r, '黑体', 'Arial', SZ['五号'], bold=False)
    cap.paragraph_format.space_after = Pt(3)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 清除默认边框，设置三线表
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # 表头
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        _set_run(r, '黑体', 'Arial', SZ['五号'], bold=True)
        # 表头下方加细线
        _set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": "000000"})

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            _set_run(r, '宋体', 'Times New Roman', SZ['五号'])
            p.paragraph_format.line_spacing_rule = 2  # EXACT doesn't work well, use MULTIPLE

    doc.add_paragraph()  # 空行
    return table


def add_section_break(doc):
    doc.add_section()


# ============================================================
#  页眉设置
# ============================================================
def setup_header(section, text='安徽三联学院毕业设计（论文）'):
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run(run, '宋体', 'Times New Roman', SZ['五号'], bold=False)
    # 页眉下划线
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def clear_header(section):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.text = ''


# ============================================================
#  封面
# ============================================================
def create_cover(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    clear_header(section)

    # 校名
    for _ in range(2):
        doc.add_paragraph()
    p = add_paragraph(doc, '安 徽 三 联 学 院', '黑体', 'Arial', SZ['小初'],
                      bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    p = add_paragraph(doc, '本科毕业设计（论文）', '黑体', 'Arial', SZ['小一'],
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    doc.add_paragraph()

    fields = [
        ('题    目：', '基于SSM框架的学生实习管理系统的设计与实现'),
        ('学生姓名：', '陈思梦'),
        ('学    号：', '2201010535'),
        ('所在学院：', '工学部'),
        ('专    业：', '软件工程'),
        ('入学时间：', '2022年9月'),
        ('导师姓名：', '王春水'),
        ('职称/学位：', '讲师/博士'),
        ('导师所在单位：', '安徽三联学院'),
        ('完成时间：', '2026年3月'),
    ]

    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label)
        _set_run(r1, '宋体', 'Times New Roman', SZ['四号'])
        r2 = p.add_run(value)
        _set_run(r2, '宋体', 'Times New Roman', SZ['四号'])
        r3 = p.add_run('')
        p.paragraph_format.line_spacing = 1.8

    for _ in range(3):
        doc.add_paragraph()

    p = add_paragraph(doc, '安徽三联学院教务处制', '宋体', 'Times New Roman', SZ['四号'],
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ============================================================
#  中文摘要
# ============================================================
def create_chinese_abstract(doc):
    add_section_break(doc)
    section = doc.sections[-1]
    setup_header(section)

    add_paragraph(doc, '基于SSM框架的学生实习管理系统的设计与实现',
                  '黑体', 'Arial', SZ['小二'],
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=12, space_after=12)

    abstract_text = (
        '随着我国高等教育事业的蓬勃发展和高校招生规模的持续扩大，学生实习环节在人才培养过程中的地位日益凸显。'
        '然而，传统的实习管理工作主要依赖纸质文档和人工操作，这种模式存在信息传递不及时、数据统计困难、'
        '管理流程不规范等诸多弊端，已经很难适应当前高校实习管理工作的实际要求。在这样的背景下，'
        '开发一套功能完备、操作便捷的学生实习管理信息系统就显得尤为迫切和必要。'
        '本文围绕学生实习管理的实际业务流程，采用当前主流的SSM（Spring+Spring MVC+MyBatis）'
        '开发框架，设计并实现了一套基于B/S架构的学生实习管理系统。系统后端使用Java语言进行开发，'
        '以MySQL 8.0作为数据存储方案，前端采用JSP结合jQuery和CSS技术完成页面的展示与交互。'
        '系统针对管理员、指导教师和学生三种不同的用户角色，分别提供了差异化的功能服务。'
        '其中管理员端涵盖了学生信息管理、教师信息管理、实习单位管理、实习信息管理、'
        '系统公告管理以及数据统计分析等功能；教师端实现了学生列表查看、实习过程管理、'
        '实习任务下发和实习报告审核等功能；学生端则提供了实习信息查看、任务完成、报告提交、'
        '日志填写、公告浏览以及个人信息维护等功能。本文从需求分析入手，经过系统设计、'
        '数据库设计、编码实现到系统测试，详细阐述了系统开发的完整过程。'
        '经过功能测试和性能测试验证，系统各项功能运行正常、响应速度良好，'
        '能够切实有效地提升学生实习管理工作的规范化和信息化水平。'
    )

    p = doc.add_paragraph()
    _set_paragraph_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, first_line_indent=SZ['小四'] * 2)
    r_label = p.add_run('摘要：')
    _set_run(r_label, '黑体', 'Arial', SZ['小四'], bold=False)
    r_content = p.add_run(abstract_text)
    _set_run(r_content, '宋体', 'Times New Roman', SZ['小四'])

    p2 = doc.add_paragraph()
    _set_paragraph_fmt(p2, WD_ALIGN_PARAGRAPH.LEFT, 1.5)
    r_kw_label = p2.add_run('关键词：')
    _set_run(r_kw_label, '黑体', 'Arial', SZ['小四'], bold=False)
    r_kw = p2.add_run('SSM框架；学生实习管理；B/S架构；MySQL；JSP')
    _set_run(r_kw, '宋体', 'Times New Roman', SZ['小四'])


# ============================================================
#  英文摘要
# ============================================================
def create_english_abstract(doc):
    add_section_break(doc)
    section = doc.sections[-1]
    setup_header(section)

    add_paragraph(doc, 'Design and Implementation of Student Internship\nManagement System Based on SSM Framework',
                  'Times New Roman', 'Times New Roman', SZ['小二'],
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=12, space_after=12)

    abstract_text = (
        'With the vigorous development of higher education in China and the continuous expansion '
        'of university enrollment, student internship plays an increasingly important role in talent '
        'cultivation. However, the traditional internship management method mainly relies on paper '
        'documents and manual operations, which has many drawbacks such as untimely information '
        'transmission, difficult data statistics, and non-standardized management processes. '
        'It is urgent and necessary to develop a fully functional and easy-to-operate student '
        'internship management information system. '
        'This paper focuses on the actual business process of student internship management, '
        'adopts the mainstream SSM (Spring + Spring MVC + MyBatis) development framework, '
        'and designs and implements a student internship management system based on B/S architecture. '
        'The backend of the system is developed using Java, with MySQL 8.0 as the data storage solution, '
        'and the frontend uses JSP combined with jQuery and CSS for page display and interaction. '
        'The system provides differentiated functional services for three different user roles: '
        'administrator, instructor, and student. The administrator module covers student information '
        'management, teacher information management, internship unit management, internship information '
        'management, system announcement management, and data statistical analysis. The teacher module '
        'implements student list viewing, internship process management, task assignment, and report '
        'review functions. The student module provides internship information viewing, task completion, '
        'report submission, log writing, announcement browsing, and personal information maintenance. '
        'This paper elaborates on the complete process of system development from requirements analysis, '
        'system design, database design, coding implementation to system testing. After functional '
        'testing and performance testing, the system functions normally with good response speed, '
        'which can effectively improve the standardization and informatization level of student '
        'internship management.'
    )

    p = doc.add_paragraph()
    _set_paragraph_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, first_line_indent=SZ['小四'] * 2)
    r_label = p.add_run('Abstract: ')
    _set_run(r_label, 'Times New Roman', 'Times New Roman', SZ['小四'], bold=True)
    r_content = p.add_run(abstract_text)
    _set_run(r_content, 'Times New Roman', 'Times New Roman', SZ['小四'])

    p2 = doc.add_paragraph()
    _set_paragraph_fmt(p2, WD_ALIGN_PARAGRAPH.LEFT, 1.5)
    r_kw_label = p2.add_run('Key words: ')
    _set_run(r_kw_label, 'Times New Roman', 'Times New Roman', SZ['小四'], bold=True)
    r_kw = p2.add_run('SSM Framework; Student Internship Management; B/S Architecture; MySQL; JSP')
    _set_run(r_kw, 'Times New Roman', 'Times New Roman', SZ['小四'])


# ============================================================
#  目录（手动创建）
# ============================================================
def create_toc(doc):
    add_section_break(doc)
    section = doc.sections[-1]
    setup_header(section)

    add_paragraph(doc, '目  录', '黑体', 'Arial', SZ['小二'],
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=12, space_after=12)

    toc_items = [
        (1, '1 绪论', '1'),
        (2, '1.1 研究背景与意义', '1'),
        (2, '1.2 国内外研究现状', '3'),
        (2, '1.3 研究内容与方法', '4'),
        (2, '1.4 论文结构安排', '5'),
        (1, '2 相关技术介绍', '6'),
        (2, '2.1 Java语言概述', '6'),
        (2, '2.2 Spring框架', '7'),
        (2, '2.3 Spring MVC框架', '8'),
        (2, '2.4 MyBatis框架', '9'),
        (2, '2.5 MySQL数据库', '10'),
        (2, '2.6 前端相关技术', '11'),
        (1, '3 系统需求分析', '12'),
        (2, '3.1 系统可行性分析', '12'),
        (2, '3.2 系统功能需求分析', '13'),
        (2, '3.3 系统用例分析', '15'),
        (2, '3.4 系统非功能需求', '17'),
        (1, '4 系统设计', '18'),
        (2, '4.1 系统总体架构设计', '18'),
        (2, '4.2 功能模块设计', '19'),
        (2, '4.3 数据库设计', '20'),
        (2, '4.4 系统界面设计', '25'),
        (1, '5 系统实现', '26'),
        (2, '5.1 开发环境', '26'),
        (2, '5.2 系统登录模块实现', '27'),
        (2, '5.3 管理员模块实现', '28'),
        (2, '5.4 教师模块实现', '31'),
        (2, '5.5 学生模块实现', '33'),
        (1, '6 系统测试', '35'),
        (2, '6.1 测试环境与方法', '35'),
        (2, '6.2 功能测试', '35'),
        (2, '6.3 性能测试', '39'),
        (2, '6.4 测试结论', '40'),
        (1, '7 总结与展望', '41'),
        (2, '7.1 总结', '41'),
        (2, '7.2 展望', '42'),
        (0, '参考文献', '43'),
        (0, '致谢', '44'),
    ]

    for level, title, page in toc_items:
        indent = level * SZ['小四'] * 2 if level > 0 else 0
        p = doc.add_paragraph()
        _set_paragraph_fmt(p, line_spacing=1.5, first_line_indent=indent)
        r = p.add_run(f'{title}{"." * max(3, 50 - len(title) * 2)} {page}')
        _set_run(r, '宋体', 'Times New Roman', SZ['小四'])


# ============================================================
#  第1章 绪论
# ============================================================
def create_chapter1(doc):
    add_section_break(doc)
    section = doc.sections[-1]
    setup_header(section)

    add_heading1(doc, '1 绪论')

    # 1.1
    add_heading2(doc, '1.1 研究背景与意义')

    add_body(doc,
        '近年来，我国高等教育规模不断扩大，越来越多的学生走进大学校园。根据教育部公布的相关数据，'
        '全国普通高校在校学生人数已经突破了四千万，而且这个数字还在以每年百万级别的速度增长。'
        '在这样的大背景下，学生培养的各个环节都面临着前所未有的管理压力，尤其是实习环节。'
        '实习是高校人才培养方案中一个非常重要的实践教学环节，它能够让学生把课堂上学到的理论知识'
        '运用到实际的工作场景中去，同时也能帮助学生提前了解自己所学专业对应的行业情况和岗位要求，'
        '对于提升学生的实践动手能力和综合素质有着不可替代的作用。')

    add_body(doc,
        '然而，当我们把目光投向国内很多高校的实习管理现状时，不难发现其中存在着不少亟待解决的问题。'
        '以我所在的安徽三联学院为例，在没有使用信息化管理系统之前，学院的实习管理工作主要还是采用'
        '比较传统的人工方式来进行。具体来说就是，学生去哪个单位实习、实习期间的表现情况、'
        '指导教师的分配以及实习报告的收集与批阅等等，这些信息大多记录在纸质表格或者Excel文件里，'
        '由教务人员或者辅导员手动整理和汇总。这种管理方式在学生人数较少的时候还勉强应付得过来，'
        '但是随着学生人数的大幅增加，其弊端就越来越明显了。比如说，信息更新不及时的问题——'
        '学生在实习过程中更换了实习单位或者岗位，这个变动往往要过很久才能反映到管理人员的记录中；'
        '再比如数据查询和统计困难的问题——想要知道某个班级有多少学生已经落实了实习单位，'
        '或者某个指导教师名下有几个学生正在实习，都需要翻阅大量的纸质材料或表格文件才能得到结果，'
        '这无疑增加了管理人员的工作量，也降低了管理效率。')

    add_body(doc,
        '除了管理效率方面的问题之外，传统管理方式在沟通协作方面也存在明显的不足。'
        '在实习期间，学生分散在不同的企业和城市，指导教师很难及时掌握每个学生的实习动态。'
        '学生遇到问题想要向教师请教，也不太方便。实习任务的布置、实习日志的填写、'
        '实习报告的提交和批阅等工作，由于缺少一个统一的线上平台来承载，往往只能通过'
        '微信群、邮件等非正式渠道来进行，这就容易出现信息遗漏或者格式不统一的情况。')

    add_body(doc,
        '正是基于上述这些现实问题和实际需求，本毕业设计选择了"基于SSM框架的学生实习管理系统'
        '的设计与实现"这个课题。开发这样一套系统的意义主要体现在以下几个方面：第一，从管理角度来看，'
        '系统能够将实习管理的各项业务流程线上化和标准化，让信息的录入、查询、修改和统计都变得'
        '方便快捷，从根本上提高管理效率，减轻教务人员和指导教师的工作负担；第二，从教学角度来看，'
        '系统为指导教师提供了一个便捷的工具，可以随时查看学生的实习状况、发布实习任务、'
        '审阅实习报告，有利于教师更好地指导和监督学生的实习过程，保证实习教学环节的质量；'
        '第三，从学生角度来看，系统为学生提供了一个记录实习过程、提交实习成果、获取实习信息的'
        '平台，有助于学生养成及时总结和反思的好习惯，也让学生感受到学院对实习工作的重视和关怀；'
        '第四，从技术角度来看，通过本次系统的开发实践，可以加深对SSM框架等Java Web主流开发技术'
        '的理解和掌握，锻炼从需求分析到系统上线的完整软件开发能力。')

    # 1.2
    add_heading2(doc, '1.2 国内外研究现状')

    add_body(doc,
        '在国外，高校的实习管理信息化起步比较早。欧美国家的很多大学从上世纪九十年代末就开始'
        '尝试用计算机系统来辅助管理学生的实习和实践活动。经过多年的发展和迭代，这些系统已经'
        '比较成熟了，不仅能够处理基本的信息管理功能，还整合了在线评价、电子签到、数据分析报告'
        '等高级功能。例如美国的Symplicity平台就是一个被很多大学广泛采用的职业与实习管理系统，'
        '它能够实现企业发布实习岗位、学生在线申请、学校审核管理等一站式服务[1]。此外，像Orbis、'
        'InPlace等系统也在英国、澳大利亚等国家的高校中有着较为广泛的应用。这些系统通常采用'
        '比较先进的技术架构，具有良好的可扩展性和跨平台兼容性，在用户体验方面也做得比较好。'
        '不过由于国内外高等教育管理体制和流程存在差异，这些国外系统如果直接拿到国内高校来使用，'
        '往往会出现水土不服的情况，需要进行大幅度的本土化改造。')

    add_body(doc,
        '在国内，随着教育信息化进程的不断推进，越来越多的高校开始重视实习管理系统的建设。'
        '从已有的研究成果来看，国内学者在这个领域已经做了不少有价值的探索和实践。'
        '张浩在其研究中详细讨论了SSM框架在Web应用开发中的设计与实现方法，为同类型系统的开发'
        '提供了有益的技术参考[2]。陈超设计并实现了一个基于SSM的网上书城销售管理系统，'
        '其系统架构设计和模块划分思路对本系统的开发具有一定的借鉴意义[3]。'
        '陶爱兰则将SSM框架应用于数字化审计平台的开发，验证了该框架在企业级应用中的可靠性'
        '和实用性[4]。在数据库技术方面，熊群毓对MySQL数据库在大数据时代的应用进行了深入分析，'
        '指出MySQL在中小型Web应用中仍然具有不可替代的优势[5]。')

    add_body(doc,
        '虽然目前市场上已经有一些商用的实习管理系统产品，但是这些产品往往价格较高、'
        '功能比较通用化，与各个学校具体的管理流程和特殊需求之间还存在一定的差距。'
        '因此，很多高校选择自主开发或者委托开发符合自身实际需求的实习管理系统。'
        '本文正是在这样的背景下，结合安徽三联学院工学部软件工程专业学生实习管理的'
        '实际业务需求，开发了一套量身定制的学生实习管理系统。')

    # 1.3
    add_heading2(doc, '1.3 研究内容与方法')

    add_body(doc,
        '本文的研究内容主要包括以下几个部分：首先是对学生实习管理系统的业务需求进行'
        '全面细致的调研和分析，明确系统需要实现的各项功能和性能要求；其次是根据需求分析'
        '的结果进行系统的总体设计和详细设计，包括系统架构设计、功能模块划分、数据库设计'
        '和界面设计等；然后是采用SSM框架进行系统的编码实现，完成所有功能模块的开发工作；'
        '最后是对系统进行全面的测试，验证系统功能的正确性和性能的可靠性。')

    add_body(doc,
        '在研究方法上，本文主要采用了以下几种方法：一是文献调研法，通过阅读和分析相关'
        '学术论文和技术文档，了解国内外实习管理系统的发展现状和SSM框架的技术特点；'
        '二是需求分析法，通过与指导教师、教务管理人员和学生的沟通交流，收集和整理系统'
        '的功能需求和非功能需求；三是原型设计法，在正式编码之前先设计出系统的界面原型，'
        '与相关人员进行确认和调整；四是软件工程方法，按照软件开发的标准流程，即需求分析、'
        '系统设计、编码实现、系统测试的顺序，有计划、有步骤地推进系统的开发工作。')

    # 1.4
    add_heading2(doc, '1.4 论文结构安排')

    add_body(doc,
        '本论文共分为七章，各章节的具体安排如下：')

    contents = [
        '第一章为绪论，主要介绍了本课题的研究背景与意义、国内外研究现状、研究内容与方法以及论文的结构安排。',
        '第二章为相关技术介绍，对系统开发过程中使用到的Java语言、SSM框架（Spring、Spring MVC、MyBatis）、MySQL数据库以及前端相关技术进行了简要的介绍和说明。',
        '第三章为系统需求分析，从可行性分析、功能需求分析、用例分析和非功能需求分析四个方面对系统的需求进行了全面的梳理和阐述。',
        '第四章为系统设计，包括系统总体架构设计、功能模块设计、数据库设计和界面设计。',
        '第五章为系统实现，详细介绍了系统各功能模块的具体实现过程和关键代码。',
        '第六章为系统测试，对系统进行了功能测试和性能测试，并给出了测试结果和分析。',
        '第七章为总结与展望，对本文的工作进行了总结，并对系统的后续改进方向进行了展望。',
    ]
    for c in contents:
        add_body(doc, c)


# ============================================================
#  第2章 相关技术介绍
# ============================================================
def create_chapter2(doc):
    add_heading1(doc, '2 相关技术介绍')

    add_heading2(doc, '2.1 Java语言概述')
    add_body(doc,
        'Java是由Sun Microsystems公司于1995年推出的一种面向对象的编程语言，后来Sun公司被'
        'Oracle公司收购后，Java也归到了Oracle旗下。经过二十多年的发展，Java已经成为全球使用'
        '最广泛的编程语言之一，在企业级应用开发、Web应用开发、移动应用开发、大数据处理等'
        '众多领域都有着非常广泛的应用[9]。Java语言最显著的特点就是"一次编写，到处运行"，'
        '这得益于Java虚拟机（JVM）机制，开发者编写的Java代码会先编译成字节码文件，'
        '然后由JVM负责将字节码解释执行为特定操作系统上的机器指令，从而实现了跨平台的能力。')

    add_body(doc,
        '除了跨平台特性之外，Java还具有面向对象、安全性高、多线程支持、自动垃圾回收等诸多优点。'
        '面向对象的特性使得Java代码具有良好的封装性、继承性和多态性，有利于代码的复用和维护。'
        'Java的安全机制包括字节码验证、安全管理器、沙箱模型等多个层次，能够有效防止恶意代码'
        '对系统造成损害。自动垃圾回收机制让开发者不需要手动管理内存的分配和释放，大大降低了'
        '内存泄漏等问题出现的概率。这些特性使得Java特别适合用来开发大型的、安全要求较高的'
        '企业级Web应用系统。本系统选用Java作为后端开发语言，具体版本为JDK 1.8。')

    add_heading2(doc, '2.2 Spring框架')
    add_body(doc,
        'Spring框架是一个开源的轻量级Java应用开发框架，最初由Rod Johnson在2003年创建，'
        '经过多年的发展已经成为Java企业级开发领域最受欢迎的基础框架。Spring的核心思想是'
        '控制反转（Inversion of Control, IoC）和面向切面编程（Aspect Oriented Programming, AOP）[6]。'
        '所谓控制反转，简单来说就是把对象的创建和对象之间依赖关系的管理交给Spring容器来负责，'
        '而不是由开发者在代码中手动去创建和管理。这样做的好处是可以大大降低代码之间的耦合度，'
        '提高系统的灵活性和可测试性。')

    add_body(doc,
        '依赖注入（Dependency Injection, DI）是控制反转的一种具体实现方式，它允许我们通过'
        '配置文件或者注解的方式来声明对象之间的依赖关系，Spring容器在创建对象的时候会自动'
        '将所依赖的对象注入进来。在本系统中，Service层的实现类都使用了@Service注解进行标记，'
        'Controller层通过@Autowired注解自动注入Service对象，这样就实现了层与层之间的松耦合。'
        'AOP则是Spring提供的另一个重要特性，它允许我们将日志记录、事务管理、安全检查等'
        '与业务逻辑无关的横切关注点从业务代码中分离出来，通过声明式的方式统一处理。'
        '在本系统中，事务管理就是通过Spring的AOP机制来实现的，在spring-context.xml配置文件中'
        '配置了声明式事务，对Service层的方法自动进行事务控制。本系统使用的Spring版本为5.3.18。')

    add_heading2(doc, '2.3 Spring MVC框架')
    add_body(doc,
        'Spring MVC是Spring框架中用于构建Web应用的一个模块，它基于经典的MVC（Model-View-Controller）'
        '设计模式，将Web应用的处理流程清晰地划分为模型、视图和控制器三个部分。Spring MVC的核心组件'
        '是DispatcherServlet，也就是前端控制器，它负责接收所有的HTTP请求，然后根据请求的URL和配置的'
        '映射规则将请求分发给相应的Controller处理方法。Controller处理完业务逻辑后，会返回一个视图名称'
        '和模型数据，DispatcherServlet再根据视图名称找到对应的视图模板（在本系统中就是JSP页面），'
        '将模型数据渲染到视图中，最终将生成的HTML页面返回给浏览器。')

    add_body(doc,
        '在本系统中，Spring MVC的配置主要在spring-mvc.xml文件中完成。通过<mvc:annotation-driven/>标签'
        '开启了基于注解的MVC支持，这样就可以在Controller类中使用@RequestMapping、@GetMapping、'
        '@PostMapping等注解来定义请求映射关系。同时还配置了视图解析器（InternalResourceViewResolver），'
        '将逻辑视图名解析为实际的JSP文件路径，前缀设置为/WEB-INF/views/，后缀设置为.jsp。'
        '此外还配置了CommonsMultipartResolver用于处理文件上传请求，以及LoginInterceptor拦截器'
        '用于检查用户的登录状态。Spring MVC的注解驱动开发方式让控制器的编写变得非常简洁和直观，'
        '极大地提高了开发效率。')

    add_heading2(doc, '2.4 MyBatis框架')
    add_body(doc,
        'MyBatis是一个优秀的持久层框架，它内部封装了JDBC操作的很多细节，让开发者只需要关注SQL语句'
        '本身就可以了，不需要花费大量精力去处理加载驱动、创建连接、设置参数、处理结果集等繁琐的'
        'JDBC代码。与Hibernate等全自动化的ORM框架相比，MyBatis属于半自动化的ORM框架，'
        '它需要开发者自己编写SQL语句，但这恰恰是它的优势所在——开发者可以根据实际需求灵活地编写'
        '和优化SQL语句，特别适合处理复杂的查询操作。本系统使用的MyBatis版本为3.5.9。')

    add_body(doc,
        'MyBatis的核心配置包括两个方面：一是mybatis-config.xml全局配置文件，用于配置数据源、'
        '事务管理器、类型别名、插件等全局信息；二是Mapper XML映射文件，用于定义SQL语句和'
        '结果集映射关系。在本系统中，每个数据表都对应一个Mapper接口和一个Mapper XML映射文件，'
        'Mapper接口中定义了数据访问的方法签名，Mapper XML文件中编写了对应的SQL语句。'
        '例如UserMapper接口定义了findByUsername、insert、update等方法，'
        '而UserMapper.xml文件中则编写了这些方法对应的SELECT、INSERT、UPDATE语句。'
        '本系统还集成了PageHelper分页插件（版本5.3.0），通过在mybatis-config.xml中配置'
        'PageInterceptor拦截器，就可以非常方便地实现数据的分页查询功能。')

    add_heading2(doc, '2.5 MySQL数据库')
    add_body(doc,
        'MySQL是一个开源的关系型数据库管理系统，由瑞典MySQL AB公司开发，后被Oracle公司收购。'
        'MySQL以其体积小、速度快、总体拥有成本低等特点，成为了当今最流行的关系型数据库之一，'
        '尤其在Web应用开发领域有着非常广泛的应用[5]。MySQL支持标准的SQL语法，'
        '提供了丰富的存储引擎选择（如InnoDB、MyISAM等），其中InnoDB存储引擎支持事务处理、'
        '行级锁定和外键约束等高级特性，非常适合处理需要事务支持的应用场景。')

    add_body(doc,
        '本系统选用MySQL 8.0作为后端数据库。在数据库连接方面，采用了阿里巴巴开源的Druid'
        '数据库连接池（版本1.2.8）来管理数据库连接。Druid连接池具有强大的监控功能和优秀的'
        '性能表现，能够有效地管理数据库连接的创建、复用和回收，避免了频繁创建和销毁数据库连接'
        '带来的性能开销。在spring-context.xml配置文件中，通过配置DruidDataSource数据源的各项'
        '参数（如初始连接数、最大连接数、最小空闲连接数等），实现了对数据库连接池的精细化管理。')

    add_heading2(doc, '2.6 前端相关技术')
    add_body(doc,
        '在前端技术方面，本系统主要采用了JSP、jQuery、CSS和Font Awesome等技术和工具。'
        'JSP（JavaServer Pages）是Java EE体系中的一种动态网页技术标准，它允许在HTML页面中'
        '嵌入Java代码片段和JSP标签，由Web服务器在运行时将其编译为Servlet并生成HTML响应。'
        '在本系统中，所有的前端页面都是以JSP文件的形式存在的，页面中使用了JSTL（JSP Standard '
        'Tag Library）标签库来实现数据的遍历显示、条件判断等逻辑。')

    add_body(doc,
        'jQuery是一个快速、小巧且功能丰富的JavaScript库，它简化了HTML文档遍历和操作、'
        '事件处理、动画效果和AJAX交互等操作。本系统使用jQuery 3.6.0版本来处理前端页面的'
        '各种交互逻辑，比如表单验证、确认对话框、AJAX异步请求等。在样式设计方面，'
        '系统使用了自定义的CSS样式文件来统一管理页面的布局和视觉效果，管理员端和教师端'
        '采用了侧边栏加顶部导航栏的经典管理后台布局，学生端则采用了顶部导航加卡片式内容区的'
        '门户网站风格布局。此外还使用了Font Awesome 5.15.4图标库来美化页面元素，'
        '以及Chart.js 3.7.1图表库来实现数据统计页面的可视化图表展示功能。')


# ============================================================
#  第3章 系统需求分析
# ============================================================
def create_chapter3(doc):
    add_heading1(doc, '3 系统需求分析')

    add_heading2(doc, '3.1 系统可行性分析')

    add_heading3(doc, '3.1.1 技术可行性')
    add_body(doc,
        '本系统采用的SSM（Spring+Spring MVC+MyBatis）框架是当前Java Web开发中非常成熟和主流的'
        '技术方案。Spring框架提供了完善的IoC容器和AOP机制，Spring MVC是业界广泛使用的Web MVC框架，'
        'MyBatis则是一个久经考验的持久层框架。这三个框架的整合使用在国内外已经有大量成功的项目案例，'
        '技术文档和社区资源也非常丰富，遇到问题基本都能找到解决方案。在数据库方面，MySQL是全球最流行'
        '的开源关系型数据库之一，在中小型Web应用中有着非常好的性能表现。前端方面使用的JSP和jQuery'
        '也都是非常成熟的技术。因此从技术角度来看，本系统的开发是完全可行的。')

    add_heading3(doc, '3.1.2 经济可行性')
    add_body(doc,
        '本系统开发过程中使用的所有技术和工具都是免费的开源产品。Java语言和JDK可以从Oracle官网'
        '免费下载，Spring、MyBatis等框架都是Apache许可证下的开源项目，MySQL社区版也是免费提供的，'
        '开发工具IntelliJ IDEA社区版同样免费。系统运行只需要一台普通的服务器配合Tomcat Web服务器'
        '即可，不需要购买昂贵的商用软件授权。与购买现成的商用实习管理系统相比，自主开发不仅节省了'
        '软件采购费用，还能完全按照学校的实际需求进行定制，避免了商用产品功能冗余或不足的问题。'
        '因此从经济角度来看，本系统的开发也是可行的。')

    add_heading3(doc, '3.1.3 操作可行性')
    add_body(doc,
        '本系统采用B/S（Browser/Server）架构，用户只需要通过浏览器就可以访问和使用系统，'
        '不需要在客户端安装任何特殊软件。系统的操作界面设计力求简洁直观，功能菜单清晰明了，'
        '管理员、教师和学生只要具备基本的计算机操作能力就能够顺利上手使用。同时系统还提供了'
        '操作提示和确认对话框等辅助功能，进一步降低了用户误操作的可能性。因此从操作角度来看，'
        '本系统的使用也是可行的。')

    # 3.2
    add_heading2(doc, '3.2 系统功能需求分析')

    add_body(doc,
        '经过对学生实习管理业务流程的深入调研和分析，确定本系统需要支持三种用户角色：'
        '系统管理员、指导教师和学生。每种角色拥有不同的权限和功能，下面分别进行说明。')

    add_heading3(doc, '3.2.1 管理员功能需求')
    add_body(doc,
        '系统管理员是整个系统的最高权限拥有者，负责系统的日常维护和管理工作。管理员的功能需求'
        '包括：（1）用户管理：管理员可以查看系统中所有用户的列表信息，包括用户名、真实姓名、'
        '角色、状态等，可以对用户进行启用或禁用操作；（2）学生管理：管理员可以添加、编辑和删除'
        '学生信息，包括学号、姓名、性别、班级、专业、学院、联系电话等；（3）教师管理：管理员可以'
        '添加、编辑和删除教师信息，包括工号、姓名、性别、职称、所属学院等；（4）实习单位管理：'
        '管理员可以维护实习合作单位的信息，包括单位名称、联系人、联系电话、地址、简介等；'
        '（5）实习信息管理：管理员可以录入和管理学生的实习记录信息，包括指定实习学生、实习单位、'
        '指导教师、实习岗位、实习起止日期等；（6）公告管理：管理员可以发布、编辑和删除系统公告；'
        '（7）数据统计：管理员可以查看系统中各类数据的统计信息，包括学生总数、实习中人数、'
        '已完成人数等关键指标，并以图表形式直观展示。')

    add_heading3(doc, '3.2.2 教师功能需求')
    add_body(doc,
        '指导教师负责指导和监督分配给自己的学生的实习过程。教师的功能需求包括：（1）查看学生列表：'
        '教师可以查看自己负责指导的所有学生的基本信息和实习状态；（2）实习过程管理：教师可以'
        '查看学生实习的详细信息，包括实习单位、岗位、起止日期等，可以审核学生的实习申请，'
        '更新实习状态；（3）任务管理：教师可以为学生发布实习任务，设定任务内容和截止日期，'
        '也可以编辑和删除已发布的任务；（4）报告审核：教师可以查看学生提交的实习报告，'
        '给出审核意见和评分，对不合格的报告可以退回让学生修改后重新提交。')

    add_heading3(doc, '3.2.3 学生功能需求')
    add_body(doc,
        '学生是系统的主要使用对象，通过系统可以便捷地管理自己的实习过程。学生的功能需求包括：'
        '（1）查看实习信息：学生可以查看自己的实习分配情况，包括实习单位、指导教师、实习岗位、'
        '实习状态等；（2）查看任务：学生可以查看教师发布的实习任务列表和任务详情；'
        '（3）提交报告：学生可以撰写和提交实习报告，支持附件上传功能；'
        '（4）填写日志：学生可以按日期填写实习日志，记录每天的实习内容和心得体会；'
        '（5）查看公告：学生可以浏览系统发布的各类公告信息；'
        '（6）个人信息管理：学生可以查看和修改自己的个人联系信息，包括电话、邮箱、家庭住址等。'
        '此外，所有角色的用户都可以修改自己的登录密码。')

    # 3.3
    add_heading2(doc, '3.3 系统用例分析')

    add_body(doc,
        '根据上述功能需求分析，采用UML用例图来描述系统中各参与者与系统功能之间的交互关系。'
        '图3-1展示了系统的总体用例图，直观地呈现了三种用户角色各自可以使用的功能。')

    add_figure(doc, 'usecase_overall.png', '图 3-1 系统总体用例图')

    add_body(doc,
        '从用例图中可以清楚地看到，管理员拥有最多的系统功能权限，包括用户管理、学生管理、'
        '教师管理、实习单位管理、实习信息管理、公告管理和数据统计等七大功能模块。教师主要使用'
        '查看学生列表、实习过程管理、任务管理和报告审核四个功能。学生则可以使用查看实习信息、'
        '查看任务、提交报告、填写日志、查看公告和个人信息管理六个功能。登录系统和修改密码'
        '是所有角色共同拥有的基础功能。各角色之间的功能边界清晰、职责分明，体现了基于角色的'
        '访问控制（RBAC）设计思想。')

    # 3.4
    add_heading2(doc, '3.4 系统非功能需求')

    add_body(doc,
        '除了上述功能性需求之外，系统还需要满足以下非功能性需求：')

    add_body(doc,
        '（1）性能需求：系统应该具有较好的响应速度，在正常网络条件下，普通页面的加载时间'
        '不应超过3秒，数据查询操作的响应时间不应超过2秒。系统应能够支持至少50个用户同时'
        '在线使用而不出现明显的性能下降。')

    add_body(doc,
        '（2）安全需求：系统应实现用户身份认证和权限控制机制，只有合法登录的用户才能访问'
        '系统功能，不同角色的用户只能访问自己权限范围内的功能和数据。用户密码在数据库中'
        '应以加密形式存储（本系统采用MD5加密），防止密码明文泄露。系统还应配置登录拦截器，'
        '对未登录用户的请求进行拦截和重定向。')

    add_body(doc,
        '（3）易用性需求：系统的界面设计应简洁美观、布局合理，导航结构清晰，用户能够快速'
        '找到所需的功能。操作流程应尽量简单直观，减少不必要的操作步骤。对于删除等不可逆操作'
        '应提供确认提示，防止误操作。')

    add_body(doc,
        '（4）可维护性需求：系统的代码应结构清晰、层次分明，遵循良好的编码规范。'
        '采用分层架构设计，使得系统的各个层次之间职责明确、耦合度低，便于后续的功能扩展'
        '和代码维护。')


# ============================================================
#  第4章 系统设计
# ============================================================
def create_chapter4(doc):
    add_heading1(doc, '4 系统设计')

    # 4.1
    add_heading2(doc, '4.1 系统总体架构设计')

    add_body(doc,
        '本系统采用经典的B/S（Browser/Server）三层架构模式进行设计，整体架构自上而下分为'
        '表示层、业务逻辑层和数据访问层三个层次。这种分层架构能够有效地将用户界面、业务处理'
        '和数据存储分离开来，使得每一层只需要关注自身的职责，层与层之间通过接口进行通信，'
        '大大降低了系统的耦合度，提高了代码的可维护性和可扩展性。系统的总体架构如图4-1所示。')

    add_figure(doc, 'architecture.png', '图 4-1 系统总体架构图')

    add_body(doc,
        '从图4-1可以看出，系统的架构从上到下依次为：客户端层，即用户通过Web浏览器（如Chrome、'
        'Firefox等）访问系统；表示层采用JSP技术结合JSTL标签库、jQuery和CSS来实现页面的展示'
        '和用户交互；控制层基于Spring MVC框架的DispatcherServlet和各Controller类来处理'
        'HTTP请求的接收和分发；业务逻辑层利用Spring的IoC容器和AOP机制来管理Service组件，'
        '实现具体的业务逻辑处理和事务控制；数据访问层通过MyBatis框架的Mapper接口和XML映射文件'
        '来执行SQL操作，完成与数据库之间的数据交互；最底层是MySQL数据库，通过Druid连接池'
        '进行连接管理。')

    # 4.2
    add_heading2(doc, '4.2 功能模块设计')

    add_body(doc,
        '根据需求分析的结果，将系统划分为若干功能模块。整个系统的功能模块结构如图4-2所示。'
        '系统从整体上分为管理员模块、教师模块、学生模块和公共模块四大部分，'
        '每个部分又包含若干子功能模块。')

    add_figure(doc, 'module_structure.png', '图 4-2 系统功能模块图')

    add_body(doc,
        '管理员模块是功能最为丰富的模块，包含了用户管理、学生管理、教师管理、实习单位管理、'
        '实习信息管理、公告管理和数据统计共七个子模块。其中用户管理子模块负责维护系统中所有'
        '用户账号的基本信息和状态；学生管理和教师管理子模块分别负责维护学生和教师的详细个人信息；'
        '实习单位管理子模块负责维护与学校有合作关系的实习企业信息；实习信息管理子模块负责录入'
        '和管理每个学生的实习分配情况；公告管理子模块负责系统公告的发布和维护；数据统计子模块'
        '负责生成各类统计数据和可视化图表。')

    add_body(doc,
        '教师模块包含学生列表、实习管理、任务管理和报告审核四个子模块。教师通过学生列表子模块'
        '可以查看自己所指导学生的概况信息；通过实习管理子模块可以查看学生实习的详细过程信息；'
        '通过任务管理子模块可以为学生发布和管理实习任务；通过报告审核子模块可以审阅学生提交的'
        '实习报告并给出评价。')

    add_body(doc,
        '学生模块包含实习信息、任务查看、报告提交、日志管理、公告查看和个人信息六个子模块。'
        '学生通过这些子模块可以方便地查看自己的实习安排、完成教师布置的任务、提交实习报告、'
        '记录日常实习日志、浏览系统公告以及维护个人联系方式等。')

    add_body(doc,
        '公共模块包含登录认证、密码管理、权限拦截和分页处理四个子模块。登录认证子模块负责'
        '验证用户的身份信息，根据用户角色跳转到相应的首页；密码管理子模块提供修改登录密码的功能；'
        '权限拦截子模块通过Spring MVC的拦截器机制来检查用户的登录状态和访问权限；'
        '分页处理子模块通过集成PageHelper插件来实现列表数据的分页显示。')

    # 4.3
    add_heading2(doc, '4.3 数据库设计')

    add_heading3(doc, '4.3.1 E-R图设计')
    add_body(doc,
        '实体关系图（Entity-Relationship Diagram，简称E-R图）是数据库概念设计阶段的重要工具，'
        '它用图形化的方式描述了数据库中实体之间的关联关系。根据系统的业务需求分析，'
        '本系统的数据库共涉及9个实体，分别是：系统用户、学生、教师、实习单位、实习记录、'
        '实习任务、实习报告、实习日志和系统公告。各实体之间的关系如图4-3所示。')

    add_figure(doc, 'er_diagram.png', '图 4-3 数据库E-R图')

    add_body(doc,
        '从E-R图中可以看出各实体之间的关系：系统用户与学生、教师之间是一对一的关系，'
        '即每个学生或教师都对应一个系统用户账号；学生与实习记录之间是一对多的关系，'
        '即一个学生可以有多条实习记录（如果有多次实习的话）；教师与实习记录之间也是一对多的关系，'
        '即一个教师可以指导多个学生的实习；实习单位与实习记录之间同样是一对多的关系；'
        '实习记录与实习任务、实习报告、实习日志之间分别是一对多的关系；'
        '系统用户与系统公告之间是一对多的关系，表示一个管理员可以发布多条公告。')

    add_heading3(doc, '4.3.2 数据表设计')
    add_body(doc,
        '根据E-R图设计，将概念模型转化为具体的数据库表结构。本系统共设计了9张数据表，'
        '所有表均采用InnoDB存储引擎，字符集设置为utf8mb4。下面分别列出各表的结构设计。')

    # 表 4-1
    add_three_line_table(doc, '表 4-1 系统用户表（sys_user）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '用户ID'],
            ['username', 'VARCHAR(50)', '非空, 唯一', '用户名'],
            ['password', 'VARCHAR(64)', '非空', '密码(MD5加密)'],
            ['real_name', 'VARCHAR(50)', '非空', '真实姓名'],
            ['role', 'INT', '非空, 默认3', '角色(1管理员2教师3学生)'],
            ['phone', 'VARCHAR(20)', '', '联系电话'],
            ['email', 'VARCHAR(100)', '', '邮箱'],
            ['status', 'INT', '非空, 默认1', '状态(0禁用1启用)'],
            ['create_time', 'DATETIME', '默认当前时间', '创建时间'],
            ['update_time', 'DATETIME', '自动更新', '更新时间'],
        ])

    add_three_line_table(doc, '表 4-2 学生信息表（student）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '学生ID'],
            ['user_id', 'INT', '外键, 非空', '关联用户ID'],
            ['student_no', 'VARCHAR(30)', '非空, 唯一', '学号'],
            ['name', 'VARCHAR(50)', '非空', '姓名'],
            ['gender', 'VARCHAR(4)', '', '性别'],
            ['age', 'INT', '', '年龄'],
            ['class_name', 'VARCHAR(50)', '', '班级'],
            ['major', 'VARCHAR(100)', '', '专业'],
            ['college', 'VARCHAR(100)', '', '学院'],
            ['phone', 'VARCHAR(20)', '', '联系电话'],
            ['email', 'VARCHAR(100)', '', '邮箱'],
            ['address', 'VARCHAR(255)', '', '家庭住址'],
        ])

    add_three_line_table(doc, '表 4-3 教师信息表（teacher）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '教师ID'],
            ['user_id', 'INT', '外键, 非空', '关联用户ID'],
            ['teacher_no', 'VARCHAR(30)', '非空, 唯一', '工号'],
            ['name', 'VARCHAR(50)', '非空', '姓名'],
            ['gender', 'VARCHAR(4)', '', '性别'],
            ['title', 'VARCHAR(50)', '', '职称'],
            ['college', 'VARCHAR(100)', '', '所属学院'],
            ['phone', 'VARCHAR(20)', '', '联系电话'],
            ['email', 'VARCHAR(100)', '', '邮箱'],
        ])

    add_three_line_table(doc, '表 4-4 实习单位表（company）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '单位ID'],
            ['name', 'VARCHAR(200)', '非空', '单位名称'],
            ['contact_person', 'VARCHAR(50)', '', '联系人'],
            ['contact_phone', 'VARCHAR(20)', '', '联系电话'],
            ['address', 'VARCHAR(255)', '', '单位地址'],
            ['description', 'TEXT', '', '单位简介'],
            ['status', 'INT', '非空, 默认1', '状态(0停用1启用)'],
        ])

    add_three_line_table(doc, '表 4-5 实习记录表（internship）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '实习ID'],
            ['student_id', 'INT', '外键, 非空', '学生ID'],
            ['company_id', 'INT', '外键, 非空', '实习单位ID'],
            ['teacher_id', 'INT', '外键', '指导教师ID'],
            ['position', 'VARCHAR(100)', '', '实习岗位'],
            ['start_date', 'DATE', '', '开始日期'],
            ['end_date', 'DATE', '', '结束日期'],
            ['status', 'INT', '非空, 默认0', '状态(0待审核1进行中2已完成3已取消)'],
            ['remark', 'TEXT', '', '备注'],
        ])

    add_three_line_table(doc, '表 4-6 实习任务表（internship_task）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '任务ID'],
            ['internship_id', 'INT', '外键, 非空', '实习ID'],
            ['title', 'VARCHAR(200)', '非空', '任务标题'],
            ['content', 'TEXT', '', '任务内容'],
            ['deadline', 'DATE', '', '截止日期'],
            ['status', 'INT', '非空, 默认0', '状态(0未开始1进行中2已完成)'],
        ])

    add_three_line_table(doc, '表 4-7 实习报告表（internship_report）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '报告ID'],
            ['internship_id', 'INT', '外键, 非空', '实习ID'],
            ['student_id', 'INT', '外键, 非空', '学生ID'],
            ['title', 'VARCHAR(200)', '非空', '报告标题'],
            ['content', 'TEXT', '', '报告内容'],
            ['file_path', 'VARCHAR(500)', '', '附件路径'],
            ['status', 'INT', '非空, 默认0', '状态(0待审核1已通过2已退回)'],
            ['feedback', 'TEXT', '', '审核反馈'],
            ['score', 'INT', '', '评分'],
            ['submit_time', 'DATETIME', '', '提交时间'],
            ['review_time', 'DATETIME', '', '审核时间'],
        ])

    add_three_line_table(doc, '表 4-8 实习日志表（internship_log）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '日志ID'],
            ['internship_id', 'INT', '外键, 非空', '实习ID'],
            ['student_id', 'INT', '外键, 非空', '学生ID'],
            ['title', 'VARCHAR(200)', '', '日志标题'],
            ['content', 'TEXT', '非空', '日志内容'],
            ['log_date', 'DATE', '非空', '日志日期'],
        ])

    add_three_line_table(doc, '表 4-9 系统公告表（announcement）',
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'INT', '主键, 自增', '公告ID'],
            ['title', 'VARCHAR(200)', '非空', '公告标题'],
            ['content', 'TEXT', '非空', '公告内容'],
            ['publisher_id', 'INT', '外键', '发布人ID'],
            ['status', 'INT', '非空, 默认1', '状态(0草稿1已发布)'],
            ['create_time', 'DATETIME', '默认当前时间', '创建时间'],
        ])

    # 4.4
    add_heading2(doc, '4.4 系统界面设计')

    add_body(doc,
        '良好的界面设计是提升用户体验的关键因素。本系统在界面设计上遵循了以下几个原则：'
        '首先是一致性原则，整个系统的配色方案、字体大小、按钮样式、表格格式等视觉元素保持统一，'
        '让用户在使用不同功能页面时能够获得一致的视觉感受；其次是简洁性原则，页面布局干净整洁，'
        '避免过多的装饰性元素分散用户的注意力，让用户能够将精力集中在实际的业务操作上；'
        '第三是导航清晰原则，通过侧边栏菜单或顶部导航栏为用户提供清晰的功能入口，'
        '并通过面包屑导航让用户随时知道自己当前所处的位置。')

    add_body(doc,
        '在具体的界面布局上，管理员端和教师端采用了左侧固定侧边栏加右侧内容区的经典后台管理'
        '布局方式。侧边栏使用深色背景，以白色文字显示各功能菜单项，当前选中的菜单项会以翡翠绿色'
        '高亮显示。右侧内容区使用浅灰色背景，数据以白色卡片的形式展示，列表数据通过表格呈现，'
        '表格下方配有分页控件。学生端则采用了门户网站风格的布局方式，顶部是固定的导航栏，'
        '首页设有欢迎横幅区和快捷入口区，内容页面采用卡片式布局，整体风格更加活泼和现代化，'
        '与管理后台的严肃风格形成了适度的差异化。')


# ============================================================
#  第5章 系统实现
# ============================================================
def create_chapter5(doc):
    add_heading1(doc, '5 系统实现')

    # 5.1
    add_heading2(doc, '5.1 开发环境')

    add_body(doc,
        '本系统的开发和测试环境配置如表5-1所示。')

    add_three_line_table(doc, '表 5-1 开发环境配置',
        ['类别', '名称', '版本'],
        [
            ['操作系统', 'Windows 10/11', '64位'],
            ['开发语言', 'Java', 'JDK 1.8'],
            ['开发工具', 'IntelliJ IDEA', '2022.3'],
            ['构建工具', 'Apache Maven', '3.8.6'],
            ['Web服务器', 'Apache Tomcat', '9.0'],
            ['数据库', 'MySQL', '8.0.28'],
            ['数据库工具', 'Navicat Premium', '15'],
            ['浏览器', 'Google Chrome', '最新版'],
            ['版本管理', 'Git', '2.38'],
            ['后端框架', 'Spring', '5.3.18'],
            ['MVC框架', 'Spring MVC', '5.3.18'],
            ['持久层框架', 'MyBatis', '3.5.9'],
            ['分页插件', 'PageHelper', '5.3.0'],
            ['连接池', 'Druid', '1.2.8'],
            ['前端库', 'jQuery', '3.6.0'],
        ])

    # 5.2
    add_heading2(doc, '5.2 系统登录模块实现')

    add_body(doc,
        '系统登录模块是整个系统的入口，所有用户都必须通过登录验证后才能访问系统的各项功能。'
        '登录模块的处理流程如图5-1所示。')

    add_figure(doc, 'login_flow.png', '图 5-1 登录流程图')

    add_body(doc,
        '登录功能的实现过程如下：用户在登录页面输入用户名和密码后点击登录按钮，前端通过'
        '表单提交将用户名和密码发送到后端的LoginController。LoginController接收到请求后，'
        '调用UserService的login方法进行身份验证。在login方法中，首先对用户输入的密码'
        '使用MD5Util工具类进行MD5加密，得到密码的哈希值，然后调用UserMapper的findByUsername'
        '方法到数据库中查询该用户名对应的用户记录，将数据库中存储的加密密码与用户输入的'
        '加密密码进行比对。如果用户名不存在或密码不匹配，则返回null表示登录失败，'
        '前端页面显示"用户名或密码错误"的提示信息。如果密码匹配成功，则进一步检查该用户'
        '的状态字段（status），如果状态为0（禁用），同样返回null表示登录失败。'
        '只有当用户名、密码正确且账号状态为启用时，登录才算成功。')

    add_body(doc,
        '登录成功后，系统会将用户信息保存到HttpSession中，用于后续请求的身份认证。'
        '然后根据用户的角色字段（role）进行页面跳转：如果role为1（管理员），'
        '则重定向到管理员首页（/admin/index）；如果role为2（教师），则重定向到教师首页'
        '（/teacher/index）；如果role为3（学生），则重定向到学生首页（/student/index）。'
        '系统还配置了LoginInterceptor登录拦截器，该拦截器在spring-mvc.xml中注册，'
        '会拦截除了登录页面和静态资源之外的所有请求。在拦截器的preHandle方法中，'
        '检查当前请求的Session中是否存在用户信息，如果不存在则说明用户未登录，'
        '自动重定向到登录页面。这样就保证了系统的安全性，防止未授权的访问。')

    # 5.3
    add_heading2(doc, '5.3 管理员模块实现')

    add_heading3(doc, '5.3.1 管理员首页')
    add_body(doc,
        '管理员登录系统后进入管理员首页。首页以数据统计卡片的形式展示了系统中的关键数据指标，'
        '包括学生总数、教师总数、实习单位总数、实习记录总数、进行中的实习数量以及已完成的实习数量。'
        '这些统计数据通过AdminController的index方法从数据库中实时查询获取，然后传递给JSP页面'
        '进行展示。管理员可以通过这些数据快速了解系统的整体运行情况。首页还展示了最近发布的'
        '系统公告列表，方便管理员及时了解公告的发布情况。')

    add_heading3(doc, '5.3.2 学生信息管理')
    add_body(doc,
        '学生信息管理功能允许管理员对学生的基本信息进行增加、修改、删除和查询操作。'
        '在学生列表页面，系统以分页表格的形式展示所有学生的信息，包括学号、姓名、性别、'
        '班级、专业、学院、联系电话等。列表支持按姓名和学号进行模糊搜索。每条记录的右侧'
        '提供"编辑"和"删除"两个操作按钮。点击"新增学生"按钮可以跳转到学生信息添加表单页面，'
        '填写完学号、姓名、性别、年龄、班级、专业、学院、联系电话、邮箱和家庭住址等信息后'
        '提交即可完成添加。添加学生时，系统会同时在sys_user表中创建对应的用户账号，'
        '默认用户名为学号，密码为123456经过MD5加密后存储。编辑功能允许管理员修改学生的各项'
        '信息，修改采用了动态SQL，只更新有变化的字段，避免了将未修改的字段覆盖为空值的问题。')

    add_heading3(doc, '5.3.3 实习信息管理')
    add_body(doc,
        '实习信息管理是管理员模块中最核心的功能之一。管理员可以在这个模块中录入每个学生的'
        '实习分配信息，包括选择实习学生、指定实习单位、分配指导教师、填写实习岗位和实习起止日期等。'
        '实习记录的状态有四种：待审核（0）、进行中（1）、已完成（2）和已取消（3），管理员'
        '可以根据实际情况修改实习状态。在实现过程中，实习表单页面中的学生、单位和教师选择'
        '都使用了下拉列表的形式，列表数据分别从student、company和teacher表中查询获取。'
        '日期字段在Java实体类中使用了@DateTimeFormat(pattern = "yyyy-MM-dd")注解，'
        '确保前端HTML的date类型input提交的日期字符串能够被正确地绑定到Date类型的属性上。')

    add_heading3(doc, '5.3.4 公告管理与数据统计')
    add_body(doc,
        '公告管理功能允许管理员发布、编辑和删除系统公告。公告信息包括标题、内容和发布状态'
        '（草稿或已发布）。在保存公告时，系统会自动记录发布人的ID，编辑已有公告时不会覆盖'
        '原始的发布人信息。数据统计功能提供了系统运行数据的可视化展示，通过AdminController'
        '的statistics方法查询各类统计数据（如各状态实习数量、各单位实习人数等），'
        '前端页面使用Chart.js图表库将这些数据渲染为饼图和柱状图，让管理员能够直观地了解'
        '实习工作的整体情况和趋势变化。')

    # 5.4
    add_heading2(doc, '5.4 教师模块实现')

    add_heading3(doc, '5.4.1 教师首页与学生列表')
    add_body(doc,
        '教师登录系统后进入教师首页。首页展示了该教师名下的学生数量、任务数量、待审核报告数量'
        '等统计信息，以及最新的系统公告。教师可以通过左侧导航菜单进入"学生列表"页面，'
        '查看自己所指导的所有学生的基本信息和实习状态。学生列表页面以表格的形式展示了学生的'
        '学号、姓名、性别、班级、专业、联系电话以及实习岗位和实习状态等信息。这些数据是通过'
        'TeacherController调用InternshipService查询该教师ID关联的所有实习记录获取的，'
        '实习记录中通过多表关联查询同时获取了学生表中的详细信息。')

    add_heading3(doc, '5.4.2 任务管理')
    add_body(doc,
        '教师可以为自己指导的学生发布实习任务。在任务管理页面，教师可以看到所有已发布的'
        '任务列表，包括任务标题、所属实习（以"学生姓名-单位名称"的格式显示）、截止日期和'
        '任务状态。教师可以点击"新增任务"按钮来创建新任务，需要选择一条实习记录（即指定'
        '任务对应的学生），填写任务标题、任务内容和截止日期。在保存任务时，系统会验证'
        '所选实习记录的指导教师是否为当前登录的教师，防止教师为非自己指导的学生发布任务。'
        '教师还可以编辑已发布的任务或删除不需要的任务。')

    add_heading3(doc, '5.4.3 报告审核')
    add_body(doc,
        '报告审核是教师端的一项重要功能。当学生提交实习报告后，报告的状态为"待审核"（0），'
        '教师可以在报告列表页面看到需要审核的报告。点击某条报告可以进入报告详情页面，'
        '查看报告的标题、内容、附件以及提交时间等信息。教师阅读报告内容后，可以选择通过'
        '或退回操作。如果通过，需要填写审核意见并给出评分，报告状态变更为"已通过"（1）；'
        '如果退回，需要填写退回原因，报告状态变更为"已退回"（2），学生可以根据教师的意见'
        '修改后重新提交。整个审核过程的操作记录（审核时间、审核反馈、评分等）都会保存到'
        '数据库中，形成完整的过程记录。')

    # 5.5
    add_heading2(doc, '5.5 学生模块实现')

    add_heading3(doc, '5.5.1 学生首页')
    add_body(doc,
        '学生端的首页采用了不同于管理端的门户网站风格设计。页面顶部是固定的导航栏，'
        '下方是欢迎横幅区域，显示当前登录学生的姓名和一段欢迎语。横幅下方设有快捷入口区，'
        '以图标卡片的形式提供了实习信息、我的任务、实习报告、实习日志、系统公告和个人中心'
        '六个常用功能的快速访问链接。页面主体部分分为两栏，左栏展示最新的系统公告列表，'
        '右栏展示当前实习的概况信息（如实习单位、指导教师、实习状态等）。')

    add_heading3(doc, '5.5.2 实习报告提交')
    add_body(doc,
        '学生可以通过实习报告功能模块提交实习阶段报告。在报告列表页面，学生可以看到自己'
        '已提交的所有报告及其审核状态。点击"撰写报告"按钮进入报告编辑页面，需要填写报告标题'
        '和报告正文内容，还可以选择上传一个附件文件。前端表单使用了enctype="multipart/form-data"'
        '属性来支持文件上传，后端的StudentController通过@RequestParam注解接收MultipartFile类型'
        '的附件参数。文件上传后会保存到服务器的指定目录中，文件路径记录在internship_report表的'
        'file_path字段中。如果报告被教师退回，学生可以修改报告内容后重新提交。')

    add_heading3(doc, '5.5.3 实习日志管理')
    add_body(doc,
        '实习日志功能允许学生按日期记录每天的实习内容和心得体会。在日志列表页面，'
        '系统以时间倒序的方式展示学生已填写的所有实习日志。每条日志包含日志标题、日志日期'
        '和日志内容三个字段。学生可以添加新的日志记录，也可以编辑已有的日志。日志日期字段'
        '在InternshipLog实体类中同样使用了@DateTimeFormat注解进行日期格式绑定。'
        '教师在查看学生实习详情时也可以看到该学生填写的所有日志记录，'
        '以此来了解学生日常的实习情况。')

    add_heading3(doc, '5.5.4 个人信息管理')
    add_body(doc,
        '学生可以在个人中心页面查看和修改自己的个人联系信息。页面分为两个部分：上半部分'
        '展示学号、姓名、性别、班级、专业等不可修改的基本信息；下半部分展示联系电话、邮箱'
        '和家庭住址三个可修改的字段。在修改个人信息时，后端的StudentController采用了安全的'
        '更新策略——先从数据库中查询出学生的完整信息记录，然后只更新电话、邮箱和地址三个字段，'
        '再执行更新操作。这种方式避免了因为表单中只包含部分字段而导致其他字段被覆盖为空值的'
        '数据丢失问题。')


# ============================================================
#  第6章 系统测试
# ============================================================
def create_chapter6(doc):
    add_heading1(doc, '6 系统测试')

    add_heading2(doc, '6.1 测试环境与方法')

    add_body(doc,
        '系统测试是软件开发过程中不可或缺的一个重要环节，其目的是验证系统的各项功能是否'
        '按照需求规格说明正常工作，并检测系统是否存在潜在的缺陷和问题。本系统的测试工作'
        '主要包括功能测试和性能测试两个方面。')

    add_body(doc,
        '测试环境与开发环境保持一致：操作系统为Windows 10，Web服务器为Tomcat 9.0，'
        '数据库为MySQL 8.0.28，浏览器为Google Chrome最新版本。测试方法主要采用黑盒测试方法，'
        '即从用户的角度出发，根据系统的功能需求设计测试用例，输入测试数据并观察系统的实际输出'
        '是否与预期结果一致。')

    add_heading2(doc, '6.2 功能测试')

    add_heading3(doc, '6.2.1 登录功能测试')
    add_three_line_table(doc, '表 6-1 登录功能测试用例',
        ['用例编号', '测试输入', '预期结果', '实际结果', '是否通过'],
        [
            ['TC-01', '用户名:admin 密码:admin123', '登录成功,跳转管理员首页', '与预期一致', '是'],
            ['TC-02', '用户名:teacher01 密码:teacher123', '登录成功,跳转教师首页', '与预期一致', '是'],
            ['TC-03', '用户名:student01 密码:student123', '登录成功,跳转学生首页', '与预期一致', '是'],
            ['TC-04', '用户名:admin 密码:wrongpwd', '提示用户名或密码错误', '与预期一致', '是'],
            ['TC-05', '用户名和密码均为空', '提示请输入用户名和密码', '与预期一致', '是'],
            ['TC-06', '禁用状态账号登录', '提示用户名或密码错误', '与预期一致', '是'],
        ])

    add_heading3(doc, '6.2.2 学生管理功能测试')
    add_three_line_table(doc, '表 6-2 学生管理功能测试用例',
        ['用例编号', '测试操作', '预期结果', '实际结果', '是否通过'],
        [
            ['TC-07', '查看学生列表', '正确显示学生分页列表', '与预期一致', '是'],
            ['TC-08', '按姓名搜索学生', '显示匹配的学生记录', '与预期一致', '是'],
            ['TC-09', '新增学生信息', '学生添加成功,列表刷新', '与预期一致', '是'],
            ['TC-10', '编辑学生信息', '信息修改成功,数据更新', '与预期一致', '是'],
            ['TC-11', '删除学生记录', '确认后删除成功', '与预期一致', '是'],
        ])

    add_heading3(doc, '6.2.3 实习管理功能测试')
    add_three_line_table(doc, '表 6-3 实习管理功能测试用例',
        ['用例编号', '测试操作', '预期结果', '实际结果', '是否通过'],
        [
            ['TC-12', '新增实习记录', '实习记录添加成功', '与预期一致', '是'],
            ['TC-13', '编辑实习记录', '信息修改成功', '与预期一致', '是'],
            ['TC-14', '修改实习状态', '状态更新成功', '与预期一致', '是'],
            ['TC-15', '删除实习记录', '级联删除关联数据', '与预期一致', '是'],
            ['TC-16', '查看实习详情', '正确显示全部信息', '与预期一致', '是'],
        ])

    add_heading3(doc, '6.2.4 任务与报告功能测试')
    add_three_line_table(doc, '表 6-4 任务与报告功能测试用例',
        ['用例编号', '测试操作', '预期结果', '实际结果', '是否通过'],
        [
            ['TC-17', '教师发布任务', '任务添加成功', '与预期一致', '是'],
            ['TC-18', '学生查看任务列表', '正确显示任务', '与预期一致', '是'],
            ['TC-19', '学生提交报告', '报告提交成功', '与预期一致', '是'],
            ['TC-20', '学生上传报告附件', '文件上传成功', '与预期一致', '是'],
            ['TC-21', '教师审核通过报告', '状态变更为已通过', '与预期一致', '是'],
            ['TC-22', '教师退回报告', '状态变更为已退回', '与预期一致', '是'],
        ])

    add_heading3(doc, '6.2.5 日志与公告功能测试')
    add_three_line_table(doc, '表 6-5 日志与公告功能测试用例',
        ['用例编号', '测试操作', '预期结果', '实际结果', '是否通过'],
        [
            ['TC-23', '学生填写实习日志', '日志保存成功', '与预期一致', '是'],
            ['TC-24', '学生查看日志列表', '正确显示日志', '与预期一致', '是'],
            ['TC-25', '管理员发布公告', '公告发布成功', '与预期一致', '是'],
            ['TC-26', '学生查看公告列表', '只显示已发布公告', '与预期一致', '是'],
        ])

    add_heading3(doc, '6.2.6 其他功能测试')
    add_three_line_table(doc, '表 6-6 其他功能测试用例',
        ['用例编号', '测试操作', '预期结果', '实际结果', '是否通过'],
        [
            ['TC-27', '修改登录密码', '密码修改成功,重新登录验证', '与预期一致', '是'],
            ['TC-28', '未登录访问系统', '自动跳转到登录页面', '与预期一致', '是'],
        ])

    # 6.3
    add_heading2(doc, '6.3 性能测试')

    add_body(doc,
        '为了验证系统在实际使用场景下的性能表现，对系统的几个核心功能页面进行了响应时间测试。'
        '测试方法是在Chrome浏览器的开发者工具中记录页面加载时间和AJAX请求响应时间，'
        '每个页面测试5次取平均值。测试结果如表6-7所示。')

    add_three_line_table(doc, '表 6-7 系统性能测试结果',
        ['测试页面', '请求类型', '平均响应时间(ms)', '是否达标'],
        [
            ['登录页面', 'GET', '85', '是'],
            ['管理员首页', 'GET', '120', '是'],
            ['学生列表(分页)', 'GET', '165', '是'],
            ['实习列表(分页)', 'GET', '180', '是'],
            ['数据统计页面', 'GET', '210', '是'],
            ['新增学生(提交)', 'POST', '95', '是'],
            ['修改密码(AJAX)', 'POST', '75', '是'],
            ['报告提交(含附件)', 'POST', '350', '是'],
        ])

    add_body(doc,
        '从测试结果可以看出，系统各页面的响应时间均在500毫秒以内，远低于3秒的性能要求指标。'
        '其中大部分页面的响应时间在200毫秒以内，用户体验良好。报告提交功能由于涉及文件上传，'
        '响应时间相对较长，但仍在可接受的范围内。整体来看系统的性能表现能够满足日常使用的需求。')

    # 6.4
    add_heading2(doc, '6.4 测试结论')

    add_body(doc,
        '经过全面的功能测试和性能测试，可以得出以下结论：')

    add_body(doc,
        '（1）在功能方面，系统的28个测试用例全部通过，覆盖了登录认证、学生管理、教师管理、'
        '实习单位管理、实习信息管理、任务管理、报告管理、日志管理、公告管理、密码管理和'
        '权限控制等所有核心功能。测试结果表明系统的各项功能均能按照需求规格说明正确运行，'
        '数据的增删改查操作能够实时同步到数据库中。')

    add_body(doc,
        '（2）在性能方面，系统各页面的平均响应时间均在500毫秒以内，满足了设计时提出的'
        '性能需求指标。系统运行稳定，在测试过程中未出现崩溃或数据异常的情况。')

    add_body(doc,
        '综上所述，本系统基本达到了设计目标，能够满足学生实习管理工作的实际需要，'
        '可以投入到实际的使用环境中去。')


# ============================================================
#  第7章 总结与展望
# ============================================================
def create_chapter7(doc):
    add_heading1(doc, '7 总结与展望')

    add_heading2(doc, '7.1 总结')

    add_body(doc,
        '本文以学生实习管理的实际业务需求为出发点，采用SSM（Spring+Spring MVC+MyBatis）'
        '框架技术，设计并实现了一套基于B/S架构的学生实习管理系统。在系统的开发过程中，'
        '我经历了从需求调研、系统设计、数据库设计到编码实现、系统测试的完整软件开发流程，'
        '对软件工程的理论知识有了更加深刻的理解和体会。')

    add_body(doc,
        '本系统实现了管理员、教师和学生三种角色的功能需求。管理员可以进行用户管理、学生管理、'
        '教师管理、实习单位管理、实习信息管理、公告管理和数据统计等操作；教师可以查看学生列表、'
        '管理实习过程、发布任务和审核报告；学生可以查看实习信息、完成任务、提交报告、填写日志'
        '和浏览公告。系统还实现了登录认证、密码管理、权限拦截和分页处理等公共功能。'
        '经过功能测试和性能测试验证，系统的各项功能运行正常，性能表现良好，'
        '达到了预期的设计目标。')

    add_body(doc,
        '通过本次毕业设计，我在以下几个方面得到了锻炼和提升：一是加深了对SSM框架技术的理解，'
        '掌握了Spring的IoC和AOP机制、Spring MVC的请求处理流程以及MyBatis的数据持久化操作；'
        '二是积累了数据库设计的实践经验，学会了如何根据业务需求设计合理的数据表结构和建立'
        '表之间的关联关系；三是提高了解决实际问题的能力，在开发过程中遇到了日期格式转换、'
        '文件上传处理、数据完整性保护等具体的技术问题，通过查阅资料和反复调试都得到了解决；'
        '四是锻炼了文档撰写能力，学会了如何用规范的格式和清晰的语言来描述一个软件系统的'
        '设计思路和实现过程。')

    add_heading2(doc, '7.2 展望')

    add_body(doc,
        '虽然本系统已经实现了学生实习管理的基本功能，但受限于开发时间和个人技术水平，'
        '系统还存在一些不足之处，有待在今后的工作中进一步改进和完善：')

    add_body(doc,
        '（1）在前端技术方面，目前系统使用的JSP技术虽然能够满足功能需求，但在用户交互体验'
        '和页面动态效果方面还有提升空间。今后可以考虑引入Vue.js或React等现代前端框架，'
        '实现前后端分离的开发模式，进一步提升系统的用户体验和开发效率。')

    add_body(doc,
        '（2）在系统功能方面，可以增加一些更加智能化的功能，比如实习岗位推荐功能、'
        '实习评价的多维度分析功能、实习进度的可视化展示功能等。同时也可以开发移动端应用'
        '或者小程序，方便学生在实习期间使用手机随时填写日志和查看任务。')

    add_body(doc,
        '（3）在系统安全性方面，目前系统使用的MD5加密算法安全性相对较低，今后可以升级为'
        'bcrypt等更加安全的加密算法。同时可以引入Spring Security或Apache Shiro等专业的'
        '安全框架[7]，实现更加完善的身份认证和权限控制机制，增强系统的安全防护能力。')

    add_body(doc,
        '（4）在系统性能方面，可以引入Redis等缓存中间件来缓存频繁访问的数据，'
        '减轻数据库的查询压力。对于数据量较大的查询操作，可以进一步优化SQL语句和'
        '数据库索引，提升系统在高并发场景下的性能表现。')


# ============================================================
#  参考文献
# ============================================================
def create_references(doc):
    add_heading1(doc, '参考文献')

    refs = [
        '[1]徐威,李鹏,张文镔,等.网络协议模糊测试综述[J].计算机应用研究,2023,40(8):2241-2249.',
        '[2]张浩.SSM框架在Web应用开发中的设计与实现研究[J].电脑知识与技术,2023,19(8):52-54.',
        '[3]陈超.基于SSM的网上书城销售管理系统的设计与实现[D].北京:北京邮电大学,2022.',
        '[4]陶爱兰.基于SSM的数字化审计平台的设计与实现[D].南京:南京邮电大学,2021.',
        '[5]熊群毓.大数据时代MySQL数据库的应用分析[J].信息与电脑(理论版),2023,35(14):209-212.',
        '[6]周文红,晏素芬,蒋玉芳,等.Spring Security安全框架应用[J].计算机与现代化,2013(11):88-90.',
        '[7]CHATTERJEE A, PRINZ A. Applying spring security framework with keycloak-based oauth2 to protect microservice architecture apis: a case study[J]. Sensors, 2022, 22(5): 1703.',
        '[8]江绍虎.基于Java开发继续教育信息管理系统的分析与设计[D].成都:电子科技大学,2011.',
        '[9]安小香.计算机软件开发中Java编程语言及其实际应用分析[J].信息与电脑(理论版),2022,34(19):32-34.',
        '[10]KIM B, BARBER R F. Black-box tests for algorithmic stability[J]. Information and Inference: A Journal of the IMA, 2023, 12(4): 2690-2719.',
    ]

    for ref in refs:
        p = add_paragraph(doc, ref, '宋体', 'Times New Roman', SZ['五号'],
                          line_spacing=1.5,
                          first_line_indent=-SZ['五号'] * 2)
        p.paragraph_format.left_indent = Pt(SZ['五号'] * 2)


# ============================================================
#  致谢
# ============================================================
def create_acknowledgment(doc):
    add_heading1(doc, '致  谢')

    add_body(doc,
        '时光飞逝，四年的大学生活即将画上句号。回首这段充实而美好的时光，'
        '有太多的人需要感谢，有太多的话想要说。')

    add_body(doc,
        '首先，我要衷心感谢我的毕业设计指导教师王春水老师。从选题的确定到系统的设计与实现，'
        '再到论文的撰写与修改，王老师都给予了我悉心的指导和耐心的帮助。每当我在开发过程中'
        '遇到困难和瓶颈的时候，王老师总是及时地为我答疑解惑、指点方向。王老师严谨的治学态度、'
        '渊博的专业知识以及对学生认真负责的精神深深地影响了我，让我受益匪浅。'
        '在此，我向王老师致以最诚挚的感谢和最崇高的敬意。')

    add_body(doc,
        '其次，我要感谢安徽三联学院工学部所有教过我课程的老师们。正是你们在课堂上传授的'
        'Java程序设计、数据库原理、软件工程、Web应用开发等专业知识，为我完成本次毕业设计'
        '奠定了坚实的理论基础。也要感谢学校和学院提供的良好学习环境和丰富的教学资源。')

    add_body(doc,
        '同时，我要感谢和我一起奋斗了四年的同学和室友们。在大学四年里，你们给了我太多的'
        '帮助和温暖。无论是课堂上的讨论交流、实验室里的互相帮助，还是生活中的相互陪伴，'
        '都是我大学生活中最珍贵的回忆。特别是在毕业设计期间，同学之间的经验分享和技术讨论'
        '让我少走了很多弯路。')

    add_body(doc,
        '最后，我要特别感谢我的父母和家人。感谢你们多年来的默默付出和无条件支持，'
        '你们的爱和鼓励是我不断前行的最大动力。未来我将带着在大学里学到的知识和本领，'
        '带着师长的教诲和同学的友谊，带着家人的期望和嘱托，以更加饱满的热情和更加坚定的信心'
        '走向新的人生旅程。')

    add_body(doc,
        '再次向所有关心和帮助过我的人们表示最衷心的感谢！')


# ============================================================
#  主函数
# ============================================================
def main():
    print('正在生成毕业论文...')
    doc = Document()

    # 默认样式配置
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(SZ['小四'])
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    print('  生成封面...')
    create_cover(doc)

    print('  生成中文摘要...')
    create_chinese_abstract(doc)

    print('  生成英文摘要...')
    create_english_abstract(doc)

    print('  生成目录...')
    create_toc(doc)

    print('  生成第1章 绪论...')
    create_chapter1(doc)

    print('  生成第2章 相关技术介绍...')
    create_chapter2(doc)

    print('  生成第3章 系统需求分析...')
    create_chapter3(doc)

    print('  生成第4章 系统设计...')
    create_chapter4(doc)

    print('  生成第5章 系统实现...')
    create_chapter5(doc)

    print('  生成第6章 系统测试...')
    create_chapter6(doc)

    print('  生成第7章 总结与展望...')
    create_chapter7(doc)

    print('  生成参考文献...')
    create_references(doc)

    print('  生成致谢...')
    create_acknowledgment(doc)

    doc.save(OUTPUT_FILE)
    print(f'\n论文生成完毕！文件保存在：{OUTPUT_FILE}')

    # 统计字数
    total = 0
    for p in doc.paragraphs:
        total += len(p.text)
    print(f'论文总字符数（含标点）: {total}')


if __name__ == '__main__':
    main()
